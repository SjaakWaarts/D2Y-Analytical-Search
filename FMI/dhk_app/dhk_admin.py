"""
Definition of views.
"""

from datetime import datetime
import re
import sys
import os
import shutil
import json
import urllib
import requests
from slugify import slugify
from pandas import Series, DataFrame
from io import BytesIO
import zipfile
import docx
from docx.table import _Cell, Table
from django.shortcuts import render
from django.http import HttpRequest
from django.http import HttpResponse
from django.http import HttpResponseRedirect
from django.template import RequestContext
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.csrf import requires_csrf_token
from django.views.decorators.csrf import ensure_csrf_cookie
import seeker.esm as esm
import FMI.settings
from FMI.settings import BASE_DIR, ES_HOSTS


def delete_recipe(request):
    # set breakpoint AFTER reading the request.body. The debugger will otherwise already consume the stream!
    json_data = json.loads(request.body)
    id = json_data.get('id', None)
    es_host = ES_HOSTS[0]
    s, search_q = esm.setup_search()
    search_filters = search_q["query"]["bool"]["filter"]
    field = 'id.keyword'
    terms = [id]
    terms_filter = {"terms": {field: terms}}
    search_filters.append(terms_filter)
    search_aggs = search_q["aggs"]
    results = esm.delete_by_query(es_host, 'recipes', search_q)
    return HttpResponse( {'succes': 'Bestand succesvol geüpload'}, content_type='application/json')

def dhk_admin_view(request):
    """Renders dhk page."""
    if request.method == 'GET':
        return render(
            request,
            'dhk_app/dhk_admin.html',
            {'site' : FMI.settings.site, 'year':datetime.now().year}
        )
    else:
        return HttpResponse({'status' : 'OK'}, content_type='application/json')

types_map = {
    '.a': 'application/octet-stream',
    '.ai': 'application/postscript',
    '.aif': 'audio/x-aiff',
    '.aifc': 'audio/x-aiff',
    '.aiff': 'audio/x-aiff',
    '.au': 'audio/basic',
    '.avi': 'video/x-msvideo',
    '.bat': 'text/plain',
    '.bcpio': 'application/x-bcpio',
    '.bin': 'application/octet-stream',
    '.bmp': 'image/bmp',
    '.c': 'text/plain',
    '.cdf': 'application/x-netcdf',
    '.cpio': 'application/x-cpio',
    '.csh': 'application/x-csh',
    '.css': 'text/css',
    '.csv': 'text/csv',
    '.dll': 'application/octet-stream',
    '.doc': 'application/msword',
    '.dot': 'application/msword',
    '.dvi': 'application/x-dvi',
    '.eml': 'message/rfc822',
    '.eps': 'application/postscript',
    '.etx': 'text/x-setext',
    '.exe': 'application/octet-stream',
    '.gif': 'image/gif',
    '.gtar': 'application/x-gtar',
    '.h': 'text/plain',
    '.hdf': 'application/x-hdf',
    '.htm': 'text/html',
    '.html': 'text/html',
    '.ico': 'image/vnd.microsoft.icon',
    '.ief': 'image/ief',
    '.jpe': 'image/jpeg',
    '.jpeg': 'image/jpeg',
    '.jpg': 'image/jpeg',
    '.js': 'application/javascript',
    '.json': 'application/json',
    '.ksh': 'text/plain',
    '.latex': 'application/x-latex',
    '.m1v': 'video/mpeg',
    '.m3u': 'application/vnd.apple.mpegurl',
    '.m3u8': 'application/vnd.apple.mpegurl',
    '.man': 'application/x-troff-man',
    '.me': 'application/x-troff-me',
    '.mht': 'message/rfc822',
    '.mhtml': 'message/rfc822',
    '.mif': 'application/x-mif',
    '.mjs': 'application/javascript',
    '.mov': 'video/quicktime',
    '.movie': 'video/x-sgi-movie',
    '.mp2': 'audio/mpeg',
    '.mp3': 'audio/mpeg',
    '.mp4': 'video/mp4',
    '.mpa': 'video/mpeg',
    '.mpe': 'video/mpeg',
    '.mpeg': 'video/mpeg',
    '.mpg': 'video/mpeg',
    '.ms': 'application/x-troff-ms',
    '.nc': 'application/x-netcdf',
    '.nws': 'message/rfc822',
    '.o': 'application/octet-stream',
    '.obj': 'application/octet-stream',
    '.oda': 'application/oda',
    '.p12': 'application/x-pkcs12',
    '.p7c': 'application/pkcs7-mime',
    '.pbm': 'image/x-portable-bitmap',
    '.pdf': 'application/pdf',
    '.pfx': 'application/x-pkcs12',
    '.pgm': 'image/x-portable-graymap',
    '.pl': 'text/plain',
    '.png': 'image/png',
    '.pnm': 'image/x-portable-anymap',
    '.pot': 'application/vnd.ms-powerpoint',
    '.ppa': 'application/vnd.ms-powerpoint',
    '.ppm': 'image/x-portable-pixmap',
    '.pps': 'application/vnd.ms-powerpoint',
    '.ppt': 'application/vnd.ms-powerpoint',
    '.ps': 'application/postscript',
    '.pwz': 'application/vnd.ms-powerpoint',
    '.py': 'text/x-python',
    '.pyc': 'application/x-python-code',
    '.pyo': 'application/x-python-code',
    '.qt': 'video/quicktime',
    '.ra': 'audio/x-pn-realaudio',
    '.ram': 'application/x-pn-realaudio',
    '.ras': 'image/x-cmu-raster',
    '.rdf': 'application/xml',
    '.rgb': 'image/x-rgb',
    '.roff': 'application/x-troff',
    '.rtx': 'text/richtext',
    '.sgm': 'text/x-sgml',
    '.sgml': 'text/x-sgml',
    '.sh': 'application/x-sh',
    '.shar': 'application/x-shar',
    '.snd': 'audio/basic',
    '.so': 'application/octet-stream',
    '.src': 'application/x-wais-source',
    '.sv4cpio': 'application/x-sv4cpio',
    '.sv4crc': 'application/x-sv4crc',
    '.svg': 'image/svg+xml',
    '.swf': 'application/x-shockwave-flash',
    '.t': 'application/x-troff',
    '.tar': 'application/x-tar',
    '.tcl': 'application/x-tcl',
    '.tex': 'application/x-tex',
    '.texi': 'application/x-texinfo',
    '.texinfo': 'application/x-texinfo',
    '.tif': 'image/tiff',
    '.tiff': 'image/tiff',
    '.tr': 'application/x-troff',
    '.tsv': 'text/tab-separated-values',
    '.txt': 'text/plain',
    '.ustar': 'application/x-ustar',
    '.vcf': 'text/x-vcard',
    '.wav': 'audio/x-wav',
    '.webm': 'video/webm',
    '.wiz': 'application/msword',
    '.wsdl': 'application/xml',
    '.xbm': 'image/x-xbitmap',
    '.xlb': 'application/vnd.ms-excel',
    '.xls': 'application/vnd.ms-excel',
    '.xml': 'text/xml',
    '.xpdl': 'application/xml',
    '.xpm': 'image/x-xpixmap',
    '.xsl': 'application/xml',
    '.xwd': 'image/x-xwindowdump',
    '.zip': 'application/zip',
}

studios = {
    'VU'    : {
        'address' : "De Vaertkant 201, 5171 JW Kaatsheuvel",
        'position': {'lat' : 51.661111, 'lng' : 5.038312}
        },
    'Stuurhuis' : {
        'address' : "Kerkstraat 19 B, 5253 AN Nieuwkuijk",
        'position': {'lat' : 51.693787, 'lng' : 5.176763}
        }
    }

def get_ext(content_type):
    ext = ""
    for ex, ct in types_map.items():
        if ct == content_type:
            ext = ex
            break
    return ext


# prevent CsrfViewMiddleware from reading the POST stream
#@csrf_exempt
@requires_csrf_token
def get_uploaded_files(request):
    # set breakpoint AFTER reading the request.body. The debugger will otherwise already consume the stream!
    json_data = json.loads(request.body)
    filter_facets = json_data.get('filter_facets', None)
    sort_facets = json_data.get('sort_facets', None)
    pager = json_data.get('pager', None)
    q = ""
    es_host = ES_HOSTS[0]
    s, search_q = esm.setup_search()
    ##
    # Workbook
    ##
    workbook = {
        'top_down_mode': 'tdn',  # td0 (flat), td1 (top-down), tdn (complete hier)
        'facets': {
            'author': {
                'label': None, 'field': 'author', 'has_keyword': True, 'size' : "1", 'multiple' : False,
                'nested': None, 'type': 'terms', 'value' : None},
            'published_date': {
                'label': None, 'field': 'published_date', 'has_keyword': False,
                'nested': None, 'type': 'date', 'value' : None},
            'title': {
                'label': None, 'field': 'title', 'has_keyword': False,
                'nested': None, 'type': 'text', 'value' : None},
        },
        'filters': ['author', 'published_date', 'title'],
        'charts': {},
        'columns': ['author', 'title', 'published_date'],
        'pager': {'nr_hits': 0, 'page_size': 10},
        'sort': [],
        'table': []
    }
    ##
    # Add Aggs
    ##
    search_aggs = search_q["aggs"]
    for facet, facet_conf in workbook['facets'].items():
        if facet in workbook['filters'] and facet_conf['type'] == 'terms':
            field = facet_conf['field']
            if facet_conf['has_keyword']:
                field = field + '.keyword'
            nested = facet_conf['nested']
            search_aggs[facet] = {'terms': {"field": field}}
            terms_agg = {'terms': {"field": field}}
            search_aggs[facet] = esm.add_agg_nesting(field, nested, terms_agg)
    ##
    # Add Filters
    ##
    if not request.user.is_staff:
        filter_facets['author'] = [request.user.username]
    search_filters = search_q["query"]["bool"]["filter"]
    search_queries = search_q["query"]["bool"]["must"]
    esm.add_search_filter(search_q, filter_facets, workbook)
    for facet_name in workbook['filters']:
        facet_conf = workbook['facets'][facet_name]
        if facet_name in filter_facets:
            if facet_conf['type'] == 'terms':
                facet_conf['value'] = filter_facets[facet_name]
            if facet_conf['type'] == 'text':
                facet_conf['value'] = filter_facets[facet_name]
            if facet_conf['type'] == 'date':
                facet_conf['value'] = filter_facets[facet_name]
            if facet_conf['type'] == 'periode':
                facet_conf['start'] = filter_facets[facet_name]['start']
                facet_conf['end'] = filter_facets[facet_name]['end']
    ##
    # Add Sorts
    ##
    if not sort_facets and workbook['sort']:
        sort_facets = workbook['sort']
    for facet, facet_conf in workbook['facets'].items():
        if facet in sort_facets:
            field = facet_conf['field']
            if facet_conf['has_keyword']:
                field = field + '.keyword'
            nested = facet_conf['nested']
            order = sort_facets[facet]
            search_q["sort"].append({field: {"order": order}})
    ##
    # Add Search
    ##
    if q is not None and q != "":
        search_q["query"]["bool"]["must"].append({"query_string": {
            "query": q, "default_operator": "AND", "fields": ["Identificatiekenmerk", "Namen"]}})
    ##
    # Add Page
    ##
    search_q["from"] = (pager['page_nr'] - 1) * pager['page_size']
    search_q["size"] = pager['page_size']
    results = esm.search_query(es_host, 'recipes', search_q)
    results = json.loads(results.text)
    hits = results.get('hits', {})
    pager['nr_hits'] = hits.get('total', {}).get('value', 0)
    aggs = results.get('aggregations', [])
    ##
    # Parse Aggs into Filter values and Charts that contain Aggs
    ##
    for facet, facet_conf in workbook['facets'].items():
        if facet in workbook['filters']:
            if facet in aggs:
                field = facet_conf['field']
                if facet_conf['has_keyword']:
                    field = field + '.keyword'
                nested = facet_conf['nested']
                buckets, totals = esm.get_buckets_nesting(field, nested, aggs[facet])
                workbook['facets'][facet]['buckets'] = buckets
    context = {
        'workbook' : workbook,
        'pager' : pager,
        'hits'  : hits,
        'aggs'  : aggs
        }
    return HttpResponse(json.dumps(context), content_type='application/json')

def table_to_df(table, header_styles, key_list, header_key_list):
    # convert the Word table to a DataFrame and/or a Dictionary
    # The DataFrame can have a header in case a certain style is used or a row first cell matches the header_key_list.
    # A Dictionary is filled when the row first cell matches the key list, those rows are not included in the DataFrame.
    # The style used by the content cells is also returned to indicate the caller on the content
    df = DataFrame()
    d = {}
    style = ""
    columns = None
    index = []
    table_row_ix = 0
    df_row_ix = 0
    matrix = []
    for table_row in table.rows[table_row_ix:]:
        header_col = table_row.cells[0]
        header_para = header_col.paragraphs[0]
        col_ix = 0
        key = header_para.text.lower()
        if key in key_list:
            d[key] = ""
            col_ix = 1
        df_row = []
        for cell in table_row.cells[col_ix:]:
            value = ""
            for para in cell.paragraphs:
                value = value + para.text
                style = para.style.name
            if key in key_list:
                d[key] = d[key] + value
            else:
                df_row.append(value)
        if key not in key_list:
            if key in header_key_list or header_para.style.name in header_styles:
                columns = df_row
            else:
                matrix.append(df_row)
                index.append(df_row_ix)
                df_row_ix = df_row_ix + 1

    df = DataFrame(matrix, columns=columns, index=index)
    return df, d, style

def has_image(par):
    """get all of the images in a paragraph 
    :param par: a paragraph object from docx
    :return: a list of r:embed 
    """
    ids = []
    root = ET.fromstring(par._p.xml)
    namespace = {
             'a':"http://schemas.openxmlformats.org/drawingml/2006/main", \
             'r':"http://schemas.openxmlformats.org/officeDocument/2006/relationships", \
             'wp':"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}

    inlines = root.findall('.//wp:inline',namespace)
    for inline in inlines:
        imgs = inline.findall('.//a:blip', namespace)
        for img in imgs:     
            id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
        ids.append(id)

def iter_block_items(parent):
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._tc
    elif isinstance(parent, docx.table._Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl):
            yield Table(child, parent)

def load_recipe(username, filename, recipe_fullname, recipe_basename, namelist):
    workshop = False
    es_host = ES_HOSTS[0]
    headers = {'Content-Type': 'application/json'}
    if 'http_auth' in es_host:
        headers['http_auth'] = es_host['http_auth']
    host = es_host['host']
    index_name = 'recipes'
    doc_type = 'recipes'
    url = "http://" + host + ":9200/" + index_name

    log = []
    recipe = {}
    recipe['id'] = recipe_basename
    recipe['title'] = os.path.splitext(filename)[0]
    #recipe['url'] = 'http://www.deheerlijkekeuken.nl/' + recipe_basename
    recipe['excerpt'] = ""
    recipe['description'] = []
    recipe['categories'] = []
    recipe['tags'] = []
    recipe['images'] = []
    recipe['cooking_clubs'] = []
    recipe['reviews'] = []
    recipe['courses'] = []

    mode = 'dish'
    doc = docx.Document(recipe_fullname)
    core_properties = doc.core_properties
    #recipe['author'] = core_properties.author
    recipe['author'] = username
    recipe['published_date'] = core_properties.created.strftime('%Y-%m-%d')

    image_type = "image"
    for shape in doc.inline_shapes:
        rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        image_part = doc.part.related_parts[rId]
        partname = image_part.partname
        if partname.startswith('/word/media/'):
            image_location = os.path.join('data', 'dhk', 'recipes', recipe_basename, 'word', 'media', partname[12:])
            image = {'type' : image_type, 'location' : image_location}
            recipe['images'].append(image)

    for block in iter_block_items(doc):
        if isinstance(block, docx.text.paragraph.Paragraph):
            para = block
            if len(para.text) > 0:
                style_name = para.style.name
                if style_name == 'Course':
                    mode = 'recipe'
                    course = {}
                    course['title'] = para.text
                    course['ingredients_parts'] = []
                    course['instructions'] = []
                    recipe['courses'].append(course)
                if mode == 'dish':
                    if style_name == 'Excerpt':
                        recipe['excerpt']= recipe['excerpt'] + para.text
                    if style_name == 'Categories':
                        pattern = re.compile("^\s+|\s*,\s*|\s+$")
                        categories_text = para.text.split(':')[1]
                        recipe['categories'].extend([x for x in pattern.split(categories_text) if x])
                    if style_name == 'Tags':
                        pattern = re.compile("^\s+|\s*,\s*|\s+$")
                        tags_text = para.text.split(':')[1]
                        recipe['tags'].extend([x for x in pattern.split(tags_text) if x])
                if mode == 'recipe':
                    if style_name == 'Recept':
                        course['instructions'].append({'instruction' : para.text})
                recipe['description'].append(para.text)
        elif isinstance(block, Table):
            table = block
            df, d, style = table_to_df(table, ['Ingredients Header'], ['studio', 'start- en eindtijd', 'kosten'], ['datum'])
            if style == 'Workshop':
                workshop = True
                for index, row in df.iterrows():
                    s = d.get('start- en eindtijd', '0:00 - 0:00').split('–')
                    start_time = s[0].strip()
                    end_time = s[1].strip()
                    cooking_club = {}
                    cooking_club['studio'] = d.get('studio','')
                    try:
                        cost_txt = d.get('kosten', '0')
                        cooking_club['cost'] = float(cost_txt)
                    except:
                        log.append("Cooking cost invalid format {0}".format(cost_txt))
                        cooking_club['cost'] = 0
                    try:
                        cooking_date_txt = row['Datum'] + " " + start_time
                        cooking_date = datetime.strptime(cooking_date_txt, '%Y-%m-%d %H:%M')
                        cooking_club['cooking_date'] = cooking_date.strftime('%Y-%m-%dT%H:%M')
                    except:
                        log.append("Cooking_date invalid format {0}".format(cooking_date_txt))
                        cooking_club['cooking_date'] = "2000-01-01T00:00"
                    cooking_club['cook'] = recipe['author']
                    cooking_club['invitation'] = row['Gerecht']
                    cooking_club['participants'] = []
                    if cooking_club['studio'] in studios:
                        for k, v in studios[cooking_club['studio']].items():
                            cooking_club[k] = v
                    recipe['cooking_clubs'].append(cooking_club)
            if style == 'Ingredients':
                for part in df.columns:
                    ingredients_part = {'part' : part, 'ingredients' : []}
                    for ingr in df[part]:
                        if len(ingr) > 0:
                            ingredients_part['ingredients'].append({'ingredient' : ingr})
                    course['ingredients_parts'].append(ingredients_part)

    success = True
    if recipe['excerpt'] == "":
        log.append("Style Excerpt is missing")
        success = False
    if len(recipe['categories']) == 0:
        log.append("Style Categories is missing")
        success = False
    if len(recipe['tags']) == 0 and not workshop:
        log.append("Style Tags is missing")
        success = False
    if len(recipe['courses']) == 0 and not workshop:
        log.append("Style Course is missing")
        success = False
    if len(recipe['images']) == 0:
        log.append("No images found")
        success = False
    for course in recipe['courses']:
        if len(course['instructions']) == 0:
            log.append("Style Recept is missing for course {0}".format(course['title']))
            success = False
        if len(course['ingredients_parts']) == 0:
            log.append("Style Ingredients is missing for course {0}".format(course['title']))
            success = False

    if success:
        data = json.dumps(recipe)
        r = requests.put(url + "/_doc/" + recipe_basename, headers=headers, data=data)
        print("load_recipe: written recipe with id", recipe_basename)
    return success, log


def ingest_recipe(username, filename, recipe_fullname):
    recipe_filename = os.path.basename(recipe_fullname)
    recipe_basename, recipe_ext = os.path.splitext(recipe_filename)
    recipe_basename = slugify(recipe_basename)
    zip_filename = recipe_basename + '.zip'
    zip_fullname = os.path.join(BASE_DIR, 'data', 'dhk', 'recipes', zip_filename)
    shutil.copy(recipe_fullname, zip_fullname)
    zip_dirname = os.path.join(BASE_DIR, 'data', 'dhk', 'recipes', recipe_basename)
    zip_ref = zipfile.ZipFile(zip_fullname, 'r')
    namelist = zip_ref.namelist()
    zip_ref.extractall(zip_dirname)
    zip_ref.close()
    success, log = load_recipe(username, filename, recipe_fullname, recipe_basename, namelist)
    os.remove(recipe_fullname)
    os.remove(zip_fullname)
    return recipe_basename, success, log

def upload_file(request):
    id = ""
    log = []
    context = {
        'id'  : id,
        'success' : False,
        'log' : log
        }
    if request.method == 'GET':
        log.append("Method is GET")
        return HttpResponse(json.dumps(context), content_type='application/json')

    try:
        request_file = request.FILES['file']
    except MultiValueDictKeyError:
        log.append("request.FILES['file'] MultiValueDictKeyError")
        return HttpResponse(json.dumps(context), content_type='application/json')

    username = request.user.username
    filename = request_file.name
    dropdown_filename = request_file.name
    dropdown_file_handler = request_file.file
    dropdown_basename, dropdown_ext = os.path.splitext(dropdown_filename)
    dropdown_basename = slugify(dropdown_basename)
    dropdown_filename = dropdown_basename + dropdown_ext
    dropdown_fullname = os.path.join(BASE_DIR, 'data', 'dhk', 'recipes', dropdown_filename)
    buffer_file = BytesIO(dropdown_file_handler.read())
    # Make sure we start copying from the beginning.
    buffer_file.seek(0)
    success = True
    try:
        with open(dropdown_fullname, 'wb') as output_file:
            shutil.copyfileobj(buffer_file, output_file)
    except IOError as e:
        id = dropdown_filename
        success = False
        log.append("Cannot open file {1}, exception {2}".format(dropdown_fullname, e))
        return HttpResponse(json.dumps(context), content_type='application/json')

    if success:
        if dropdown_ext == '.zip':
            zip_dirname = os.path.join(BASE_DIR, 'data', 'dhk', 'recipes', dropdown_basename)
            zip_ref = zipfile.ZipFile(dropdown_fullname, 'r')
            zip_ref.extractall(zip_dirname)
            zip_ref.close()
        elif dropdown_ext == '.docx':
            id, success, log = ingest_recipe(username, filename, dropdown_fullname)

    context = {
        'id'  : id,
        'success' : success,
        'log' : log
        }
    return HttpResponse(json.dumps(context), content_type='application/json')


def stream_file(request):
    location = request.GET.get('location', None)
    if sys.platform[0:3] == "win":
        location = location.replace('/', '\\')
    else:
        location = location.replace('\\', '/')
    filename = os.path.join(BASE_DIR, location)
    if os.path.isfile(filename):
        #mime = magic.Magic(mime=True)
        #content_type = mime.from_file(filename)
        splitext = os.path.splitext(filename)
        basename = os.path.basename(filename)
        content_type = types_map[splitext[1]]
        with open(filename, "rb") as f:
            response = HttpResponse(f.read(), content_type=content_type)
        #response._headers['Content-Disposition'] = "attachment; filename=" + basename
        return response
    else:
        response = HttpResponse(content_type="image/jpeg")
        return response

